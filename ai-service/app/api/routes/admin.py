import asyncio
import io
import logging
from pathlib import Path
from fastapi import APIRouter, UploadFile, File, HTTPException
from app.models.schemas import DatasetUploadResponse, DatasetListResponse
from app.services.chroma_service import chroma_service
from app.services.chunking_service import chunking_service
from app.services.llm_factory import get_embed_service

logger = logging.getLogger(__name__)

router = APIRouter()


def _ocr_pdf(content: bytes) -> str:
    try:
        from pdf2image import convert_from_bytes
        import pytesseract
        pytesseract.pytesseract.tesseract_cmd = "/opt/homebrew/bin/tesseract"
        images = convert_from_bytes(content, dpi=200)
        pages = []
        for img in images:
            text = pytesseract.image_to_string(img, lang="rus+kaz+eng")
            if text.strip():
                pages.append(text)
        result = "\n\n".join(pages)
        logger.info("OCR extracted %d chars from %d pages", len(result), len(images))
        return result
    except Exception as e:
        logger.error("OCR failed: %s", e)
        return ""


async def _extract_text(file: UploadFile) -> str:
    content = await file.read()
    name = (file.filename or "").lower()
    logger.info("_extract_text: filename=%r name=%r content_len=%d", file.filename, name, len(content))

    if name.endswith(".pdf"):
        try:
            import pypdf
            reader = pypdf.PdfReader(io.BytesIO(content))
            text = "\n\n".join(page.extract_text() or "" for page in reader.pages)
            if text.strip():
                return text
            logger.info("pypdf returned empty text, trying OCR for %r", file.filename)
            return _ocr_pdf(content)
        except Exception as e:
            raise HTTPException(status_code=400, detail=f"PDF parse error: {e}")

    if name.endswith(".docx"):
        try:
            import docx
            doc = docx.Document(io.BytesIO(content))
            parts = [p.text for p in doc.paragraphs if p.text.strip()]
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            parts.append(cell.text)
            text = "\n\n".join(parts)
            logger.info("DOCX parsed: %d chars, %d paragraphs, %d tables", len(text), len(doc.paragraphs), len(doc.tables))
            return text
        except Exception as e:
            logger.error("DOCX exception: %s: %s", type(e).__name__, e)
            raise HTTPException(status_code=400, detail=f"DOCX parse error: {e}")

    if name.endswith(".doc"):
        try:
            import subprocess, tempfile, os
            with tempfile.NamedTemporaryFile(suffix=".doc", delete=False) as tmp:
                tmp.write(content)
                tmp_path = tmp.name
            result = subprocess.run(
                ["textutil", "-convert", "txt", "-stdout", tmp_path],
                capture_output=True, text=True, timeout=30
            )
            os.unlink(tmp_path)
            logger.info("textutil returncode=%s stdout_len=%s stderr=%s", result.returncode, len(result.stdout), result.stderr)
            if result.returncode != 0:
                raise HTTPException(status_code=400, detail=f"DOC parse error: {result.stderr}")
            return result.stdout
        except HTTPException:
            raise
        except Exception as e:
            logger.error("DOC exception: %s: %s", type(e).__name__, e)
            raise HTTPException(status_code=400, detail=f"DOC parse error: {e}")

    if name.endswith(".xlsx"):
        try:
            import openpyxl
            wb = openpyxl.load_workbook(io.BytesIO(content), read_only=True, data_only=True)
            parts = []
            for sheet in wb.worksheets:
                parts.append(f"[Лист: {sheet.title}]")
                for row in sheet.iter_rows(values_only=True):
                    line = "\t".join(str(c) if c is not None else "" for c in row)
                    if line.strip():
                        parts.append(line)
            return "\n".join(parts)
        except Exception as e:
            raise HTTPException(status_code=400, detail=f"XLSX parse error: {e}")

    return content.decode("utf-8", errors="ignore")


MAX_DOC_CHARS = 30_000
EMBED_BATCH_SIZE = 10

@router.post("/admin/datasets/upload", response_model=DatasetUploadResponse)
async def upload_dataset(file: UploadFile = File(...)):
    source = Path(file.filename or "dataset").stem
    text = await _extract_text(file)
    if len(text) > MAX_DOC_CHARS:
        logger.info("Document truncated from %d to %d chars", len(text), MAX_DOC_CHARS)
        text = text[:MAX_DOC_CHARS]
    chunks = chunking_service.chunk(text)
    if not chunks:
        raise HTTPException(status_code=400, detail="Файл пустой или не удалось извлечь текст")
    embed_service = get_embed_service()
    logger.info("Embedding %d chunks for source='%s'", len(chunks), source)
    embeddings = []
    for i in range(0, len(chunks), EMBED_BATCH_SIZE):
        batch = chunks[i:i + EMBED_BATCH_SIZE]
        batch_embeddings = await asyncio.gather(*[embed_service.get_embedding(c) for c in batch])
        embeddings.extend(batch_embeddings)
    chroma_service.add_static_chunks(list(chunks), embeddings, source)
    return DatasetUploadResponse(source=source, chunks_indexed=len(chunks))


@router.get("/admin/datasets", response_model=DatasetListResponse)
async def list_datasets():
    return DatasetListResponse(sources=chroma_service.list_static_sources())


@router.post("/admin/extract-text")
async def extract_text_endpoint(file: UploadFile = File(...)):
    text = await _extract_text(file)
    return {"text": text}


@router.delete("/admin/datasets/{source}")
async def delete_dataset(source: str):
    chroma_service.delete_static_source(source)
    return {"deleted": True, "source": source}
